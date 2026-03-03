"""Convert ctxpack-whitepaper-v2.md to .docx using python-docx.

High-quality conversion with:
  - Rich text in table cells (bold, italic, code)
  - Unicode footnote markers (¹ ² ³)
  - Proper heading hierarchy (subsections like 5.10.1)
  - Paragraph spacing, page numbers, narrow margins
  - Code blocks in Consolas with shading

Usage: python paper/md_to_docx.py [input.md] [output.docx]
"""

import os
import re
import sys

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def md_to_docx(md_path: str, docx_path: str) -> None:
    with open(md_path, encoding="utf-8") as f:
        text = f.read()

    doc = Document()

    # ── Document styles ──
    _setup_styles(doc)

    lines = text.split("\n")
    i = 0
    in_code_block = False
    code_buffer: list[str] = []
    in_table = False
    table_rows: list[list[str]] = []
    table_alignments: list[str] = []
    footnote_lines: list[str] = []

    while i < len(lines):
        line = lines[i]

        # ── Code blocks ──
        if line.strip().startswith("```"):
            if in_code_block:
                _add_code_block(doc, code_buffer)
                code_buffer = []
                in_code_block = False
            else:
                if in_table:
                    _add_table(doc, table_rows, table_alignments)
                    table_rows = []
                    table_alignments = []
                    in_table = False
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ── Table rows ──
        if line.strip().startswith("|") and "|" in line.strip()[1:]:
            # Separator row — extract alignments
            if re.match(r"^\s*\|[\s\-:|]+\|\s*$", line):
                if not table_alignments:
                    table_alignments = _parse_table_alignments(line)
                i += 1
                continue
            row = [cell.strip() for cell in line.strip().strip("|").split("|")]
            if not in_table:
                in_table = True
            table_rows.append(row)
            i += 1
            continue
        else:
            if in_table:
                _add_table(doc, table_rows, table_alignments)
                table_rows = []
                table_alignments = []
                in_table = False

        stripped = line.strip()

        # Blank line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if stripped == "---":
            i += 1
            continue

        # Images: ![alt](path)
        img_m = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", stripped)
        if img_m:
            alt_text = img_m.group(1)
            img_path = img_m.group(2)
            # Resolve relative to markdown file directory
            if not os.path.isabs(img_path):
                img_path = os.path.join(os.path.dirname(md_path), img_path)
            if os.path.exists(img_path):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run()
                run.add_picture(img_path, width=Inches(6.0))
                # Caption below
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_before = Pt(2)
                cap.paragraph_format.space_after = Pt(8)
                cap_run = cap.add_run(alt_text)
                cap_run.font.size = Pt(9)
                cap_run.italic = True
                cap_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            else:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(f"[Image not found: {img_m.group(2)}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            i += 1
            continue

        # Footnote lines (¹ ² ³ style) — collect for end
        if re.match(r"^[¹²³⁴⁵⁶⁷⁸⁹⁰]+\s", stripped):
            footnote_lines.append(stripped)
            i += 1
            continue

        # Headings
        m = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if m:
            level = len(m.group(1))
            heading_text = m.group(2)
            if level == 1:
                # Title
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.paragraph_format.space_before = Pt(24)
                p.paragraph_format.space_after = Pt(6)
                run = p.add_run(_clean_md(heading_text))
                run.bold = True
                run.font.size = Pt(18)
                run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
            else:
                heading_level = min(level, 4)
                h = doc.add_heading(level=heading_level)
                _add_rich_text(h, heading_text)
                # Tighter spacing for subsections
                if level >= 4:
                    h.paragraph_format.space_before = Pt(8)
                    h.paragraph_format.space_after = Pt(4)
            i += 1
            continue

        # Author line (bold-only, centred)
        if stripped.startswith("**") and stripped.endswith("**") and len(stripped) < 80:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(_clean_md(stripped))
            run.bold = True
            run.font.size = Pt(12)
            i += 1
            continue

        # Centred non-bold line after author (e.g., "Independent Researcher")
        if i > 0 and i < len(lines) - 1:
            prev = lines[i - 1].strip()
            if prev.startswith("**") and prev.endswith("**") and len(prev) < 80 and len(stripped) < 80 and not stripped.startswith("#"):
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(stripped)
                run.italic = True
                run.font.size = Pt(11)
                p.paragraph_format.space_after = Pt(12)
                i += 1
                continue

        # Bullet / numbered list
        m = re.match(r"^(\s*)([-*])\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            _add_rich_text(p, m.group(3))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        m = re.match(r"^(\s*)(\d+)\.\s+(.+)$", stripped)
        if m:
            p = doc.add_paragraph(style="List Number")
            _add_rich_text(p, m.group(3))
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            i += 1
            continue

        # LaTeX formula
        if stripped.startswith("$$"):
            formula = stripped.strip("$").strip()
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(formula)
            run.italic = True
            run.font.size = Pt(10)
            i += 1
            continue

        # Regular paragraph
        p = doc.add_paragraph()
        _add_rich_text(p, stripped)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        i += 1

    # Flush remaining table
    if in_table:
        _add_table(doc, table_rows, table_alignments)

    # Add collected footnotes
    if footnote_lines:
        # Thin rule
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        run = p.add_run("_" * 40)
        run.font.size = Pt(6)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

        for fn in footnote_lines:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            _add_rich_text(p, fn)
            for run in p.runs:
                run.font.size = Pt(8.5)

    doc.save(docx_path)
    print(f"Created: {docx_path}")


def _setup_styles(doc: Document) -> None:
    """Configure document styles: fonts, margins, page numbers."""
    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    style.paragraph_format.space_before = Pt(3)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.15

    # Heading styles
    for level in range(1, 5):
        h_name = f"Heading {level}"
        if h_name in doc.styles:
            h_style = doc.styles[h_name]
            h_style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

        # Page numbers in footer
        footer = section.footer
        footer.is_linked_to_previous = False
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _add_page_number(fp)
        for run in fp.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def _add_page_number(paragraph) -> None:
    """Insert a page number field into a paragraph."""
    run = paragraph.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fld_char_begin)

    run2 = paragraph.add_run()
    instr = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instr)

    run3 = paragraph.add_run()
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fld_char_end)


def _parse_table_alignments(sep_line: str) -> list[str]:
    """Parse column alignments from separator row like |:---:|---:|."""
    cells = sep_line.strip().strip("|").split("|")
    aligns = []
    for cell in cells:
        cell = cell.strip()
        if cell.startswith(":") and cell.endswith(":"):
            aligns.append("center")
        elif cell.endswith(":"):
            aligns.append("right")
        else:
            aligns.append("left")
    return aligns


def _add_rich_text(paragraph, text: str) -> None:
    """Add text with inline formatting (bold, italic, code, links, footnote markers)."""
    # Split on formatting markers
    parts = re.split(
        r"(\*\*[^*]+\*\*"        # bold
        r"|\*[^*]+\*"            # italic
        r"|`[^`]+`"              # inline code
        r"|\[[^\]]+\]\([^)]+\)"  # links
        r"|[¹²³⁴⁵⁶⁷⁸⁹⁰]+)"    # footnote markers
        , text
    )
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*") and not part.startswith("**"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif part.startswith("["):
            m = re.match(r"\[([^\]]+)\]\(([^)]+)\)", part)
            if m:
                run = paragraph.add_run(m.group(1))
                run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
                run.underline = True
            else:
                paragraph.add_run(part)
        elif re.match(r"^[¹²³⁴⁵⁶⁷⁸⁹⁰]+$", part):
            # Superscript footnote marker
            run = paragraph.add_run(part)
            run.font.superscript = True
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB3)
        else:
            paragraph.add_run(part)


def _add_code_block(doc, lines: list[str]) -> None:
    """Add a code block with monospace font and light grey background."""
    code_text = "\n".join(lines)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.right_indent = Inches(0.25)

    # Light grey shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5" w:val="clear"/>')
    p._p.get_or_add_pPr().append(shading)

    run = p.add_run(code_text)
    run.font.name = "Consolas"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def _add_table(doc, rows: list[list[str]], alignments: list[str]) -> None:
    """Add a table with rich text in cells, alignment, and styling."""
    if not rows:
        return

    max_cols = max(len(r) for r in rows)

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            if col_idx >= max_cols:
                continue
            cell = table.cell(row_idx, col_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)

            # Apply column alignment
            if col_idx < len(alignments):
                align = alignments[col_idx]
                if align == "center":
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif align == "right":
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

            if row_idx == 0:
                # Header row: bold
                _add_rich_text(p, cell_text)
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(8.5)
                # Header shading
                shading = parse_xml(
                    f'<w:shd {nsdecls("w")} w:fill="E8EEF4" w:val="clear"/>'
                )
                cell._tc.get_or_add_tcPr().append(shading)
            else:
                _add_rich_text(p, cell_text)
                for run in p.runs:
                    run.font.size = Pt(8.5)

    # Auto-fit column widths based on content
    total_width = 6.5  # inches available
    col_weights = []
    for col_idx in range(max_cols):
        max_len = 0
        for row_data in rows:
            if col_idx < len(row_data):
                max_len = max(max_len, len(_clean_md(row_data[col_idx])))
        col_weights.append(max(max_len, 3))
    total_weight = sum(col_weights)

    for col_idx in range(max_cols):
        width = Inches(total_width * col_weights[col_idx] / total_weight)
        for row in table.rows:
            if col_idx < len(row.cells):
                row.cells[col_idx].width = width


def _clean_md(text: str) -> str:
    """Remove markdown formatting from text."""
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"[¹²³⁴⁵⁶⁷⁸⁹⁰]+", "", text)
    return text.strip()


if __name__ == "__main__":
    script_dir = os.path.dirname(os.path.abspath(__file__))

    if len(sys.argv) >= 2:
        md_path = sys.argv[1]
    else:
        md_path = os.path.join(script_dir, "ctxpack-whitepaper-v2.md")

    if len(sys.argv) >= 3:
        docx_path = sys.argv[2]
    else:
        docx_path = os.path.join(script_dir, "ctxpack-whitepaper-v2.docx")

    if not os.path.exists(md_path):
        print(f"Error: {md_path} not found", file=sys.stderr)
        sys.exit(1)

    md_to_docx(md_path, docx_path)
